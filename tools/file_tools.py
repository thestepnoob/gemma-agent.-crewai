import os
import json
from crewai.tools import BaseTool
from .core import is_safe_path, is_protected_code_path, log_created_file, was_file_created_by_agent, AGENT_LOG_FILE

class ListDirectoryTool(BaseTool):
    name: str = "Verzeichnis auflisten"
    description: str = "Listet alle Dateien und Ordner in einem sicheren Pfad unter C:\\Users\\iboer auf."

    def _run(self, directory_path: str) -> str:
        if not is_safe_path(directory_path):
            return f"Zugriff verweigert: Der Pfad '{directory_path}' liegt außerhalb des erlaubten Verzeichnisses."
        
        try:
            if not os.path.exists(directory_path):
                return f"Fehler: Das Verzeichnis '{directory_path}' existiert nicht."
            
            items = os.listdir(directory_path)
            result = []
            for item in items:
                full_path = os.path.join(directory_path, item)
                if os.path.isdir(full_path):
                    result.append(f"[ORDNER] {item}")
                else:
                    result.append(f"[DATEI]  {item}")
            return "\n".join(result) if result else "Verzeichnis ist leer."
        except Exception as e:
            return f"Fehler beim Auflisten des Verzeichnisses: {str(e)}"

class ReadFileTool(BaseTool):
    name: str = "Datei lesen"
    description: str = "Liest den Inhalt einer Text-, Code- oder PDF-Datei."

    def _run(self, file_path: str) -> str:
        if not is_safe_path(file_path):
            return f"Zugriff verweigert: Der Pfad '{file_path}' ist nicht erlaubt."
        
        try:
            if not os.path.exists(file_path):
                return f"Fehler: Die Datei '{file_path}' existiert nicht."
            if os.path.isdir(file_path):
                return f"Fehler: '{file_path}' ist ein Verzeichnis, keine Datei."
            
            if file_path.lower().endswith(".pdf"):
                try:
                    import pdfplumber
                    text_parts = []
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                    return "\n".join(text_parts) if text_parts else "Die PDF-Datei enthält keinen extrahierbaren Text."
                except Exception as pdf_err:
                    return f"Fehler beim Extrahieren des Texts aus der PDF: {str(pdf_err)}"

            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return content
        except Exception as e:
            return f"Fehler beim Lesen der Datei: {str(e)}"

class WriteFileTool(BaseTool):
    name: str = "Datei schreiben"
    description: str = "Schreibt oder überschreibt eine Datei mit neuem Inhalt."

    def _run(self, file_path: str, content: str) -> str:
        if not is_safe_path(file_path):
            return f"Zugriff verweigert: Der Pfad '{file_path}' ist nicht erlaubt."
        if is_protected_code_path(file_path):
            return f"Zugriff verweigert: Das Modifizieren von Python-Quellcode im Bot-Verzeichnis ist untersagt."
        
        try:
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)
            
            log_created_file(file_path)
            return f"Erfolgreich: Datei '{file_path}' geschrieben und als vom Agenten erstellt registriert."
        except Exception as e:
            return f"Fehler beim Schreiben der Datei: {str(e)}"

class DeleteFileTool(BaseTool):
    name: str = "Datei loeschen"
    description: str = (
        "Löscht eine Datei. ACHTUNG: Wenn du eine Datei löschen willst, die NICHT von dir selbst erstellt wurde, "
        "musst du zwingend 'user_confirmed=True' übergeben, nachdem du den Benutzer im Chat "
        "gefragt hast und er dir die Erlaubnis erteilt hat."
    )

    def _run(self, file_path: str, user_confirmed: bool = False) -> str:
        if not is_safe_path(file_path):
            return f"Zugriff verweigert: Der Pfad '{file_path}' ist nicht erlaubt."
        if is_protected_code_path(file_path):
            return f"Zugriff verweigert: Das Löschen von Python-Quellcode im Bot-Verzeichnis ist untersagt."
        
        try:
            if not os.path.exists(file_path):
                return f"Fehler: Die Datei '{file_path}' existiert nicht."
            if os.path.isdir(file_path):
                return f"Fehler: '{file_path}' ist ein Verzeichnis. Dieses Tool löscht nur einzelne Dateien."
            
            self_created = was_file_created_by_agent(file_path)
            
            if not self_created and not user_confirmed:
                return (
                    f"LÖSCHUNG ABGELEHNT: Die Datei '{file_path}' wurde nicht von dir erstellt.\n"
                    "Bitte frage den Benutzer im Chat um Erlaubnis, diese Datei zu löschen. "
                    "Sobald zugestimmt wurde, rufe dieses Tool mit 'user_confirmed=True' auf."
                )
            
            os.remove(file_path)
            
            if self_created and os.path.exists(AGENT_LOG_FILE):
                try:
                    with open(AGENT_LOG_FILE, "r", encoding="utf-8") as f:
                        lines = f.readlines()
                    with open(AGENT_LOG_FILE, "w", encoding="utf-8") as f:
                        for line in lines:
                            if line.strip().lower() != os.path.abspath(file_path).lower():
                                f.write(line)
                except Exception:
                    pass
                    
            status = "selbst erstellte" if self_created else "fremde (benutzerbestätigte)"
            return f"Erfolgreich: Die {status} Datei '{file_path}' wurde gelöscht."
        except Exception as e:
            return f"Fehler beim Löschen der Datei: {str(e)}"

class ReadWordDocumentTool(BaseTool):
    name: str = "Word Datei lesen"
    description: str = "Liest den Textinhalt einer Word-Datei (.docx)."

    def _run(self, file_path: str) -> str:
        if not is_safe_path(file_path):
            return f"Zugriff verweigert: Pfad ist nicht erlaubt."
        
        try:
            import docx
            if not os.path.exists(file_path):
                return f"Fehler: Die Datei '{file_path}' existiert nicht."
            
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    full_text.append(" | ".join(row_text))
                    
            return "\n".join(full_text)
        except Exception as e:
            return f"Fehler beim Lesen der Word-Datei: {str(e)}"

class ReadExcelDocumentTool(BaseTool):
    name: str = "Excel oder CSV lesen"
    description: str = "Liest die Tabellendaten einer Excel-Datei (.xlsx) oder CSV-Datei (.csv)."

    def _run(self, file_path: str) -> str:
        if not is_safe_path(file_path):
            return f"Zugriff verweigert: Pfad ist nicht erlaubt."
        
        try:
            if not os.path.exists(file_path):
                return f"Fehler: Die Datei '{file_path}' existiert nicht."
            
            if file_path.lower().endswith(".csv"):
                import csv
                output = []
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    sample = f.read(2048)
                    f.seek(0)
                    dialect = csv.Sniffer().sniff(sample) if sample else ","
                    reader = csv.reader(f, dialect)
                    for row in reader:
                        output.append(" | ".join(row))
                return "\n".join(output)
            
            elif file_path.lower().endswith(".xlsx"):
                import openpyxl
                wb = openpyxl.load_workbook(file_path, data_only=True)
                output = []
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    output.append(f"--- Tabelle: {sheet_name} ---")
                    for row in sheet.iter_rows(values_only=True):
                        if any(row):
                            row_str = [str(val) if val is not None else "" for val in row]
                            output.append(" | ".join(row_str))
                return "\n".join(output)
            
            else:
                return "Fehler: Nicht unterstütztes Format. Dieses Tool liest nur .xlsx und .csv."
        except Exception as e:
            return f"Fehler beim Lesen der Excel/CSV-Datei: {str(e)}"

class WriteExcelDocumentTool(BaseTool):
    name: str = "Excel oder CSV schreiben"
    description: str = (
        "Erstellt oder überschreibt eine Excel-Datei (.xlsx) oder CSV-Datei (.csv). "
        "Die Daten müssen als JSON-Array von Arrays übergeben werden (data_json), wobei jede Zeile eine Liste von Werten darstellt. "
        "Beispiel: '[[\"Name\", \"Alter\"], [\"Max\", 25]]'"
    )

    def _run(self, file_path: str, data_json: str) -> str:
        if not is_safe_path(file_path):
            return f"Zugriff verweigert: Pfad ist nicht erlaubt."
        if is_protected_code_path(file_path):
            return f"Zugriff verweigert: Modifikation geschützter Code-Dateien ist untersagt."
        
        try:
            data = json.loads(data_json)
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            if file_path.lower().endswith(".csv"):
                import csv
                with open(file_path, "w", encoding="utf-8", newline="") as f:
                    writer = csv.writer(f)
                    writer.writerows(data)
                log_created_file(file_path)
                return f"Erfolgreich: CSV-Datei '{file_path}' wurde geschrieben."
            
            elif file_path.lower().endswith(".xlsx"):
                import openpyxl
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "Daten"
                for row in data:
                    ws.append(row)
                wb.save(file_path)
                log_created_file(file_path)
                return f"Erfolgreich: Excel-Datei '{file_path}' wurde geschrieben."
            
            else:
                return "Fehler: Nur .xlsx und .csv Dateiendungen sind erlaubt."
        except Exception as e:
            return f"Fehler beim Schreiben der Datei: {str(e)}"
